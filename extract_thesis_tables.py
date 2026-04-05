#!/usr/bin/env python3
"""
Extract specific tables and paragraph context from Tesis.docx
Tables 32-38, 40-45, plus sampling methodology paragraphs.
"""

import sys
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import re
import statistics

DOCX_PATH = "/home/royer/Desktop/project-citas-medicas/Tesis.docx"

# ──────────────────────────────────────────────────────
# Helper: extract text from a cell, including OMML math
# ──────────────────────────────────────────────────────
def extract_cell_text(cell):
    """Extract all text from a cell, including Office Math (OMML) formulas."""
    parts = []
    for paragraph in cell.paragraphs:
        p_parts = []
        # Walk through the paragraph XML to pick up both runs and math
        for child in paragraph._element:
            tag = etree.QName(child.tag).localname if '}' in child.tag else child.tag
            if tag == 'r':  # normal run
                t = child.find(qn('w:t'))
                if t is not None and t.text:
                    p_parts.append(t.text)
            elif tag == 'oMath' or tag == 'oMathPara':
                # Try to pull readable text out of the OMML XML
                math_text = _omml_to_text(child)
                if math_text:
                    p_parts.append(f"[MATH: {math_text}]")
            elif tag == 'hyperlink':
                for r in child.findall(qn('w:r')):
                    t = r.find(qn('w:t'))
                    if t is not None and t.text:
                        p_parts.append(t.text)
        # Fallback: if we got nothing from XML walk, use python-docx text
        if not p_parts and paragraph.text.strip():
            p_parts.append(paragraph.text.strip())
        if p_parts:
            parts.append("".join(p_parts))
    return "\n".join(parts)


def _omml_to_text(elem):
    """Recursively extract readable text from an OMML element."""
    texts = []
    # m:t elements hold the actual text/symbols inside OMML
    for t in elem.iter():
        local = etree.QName(t.tag).localname if '}' in t.tag else t.tag
        if local == 't' and t.text:
            texts.append(t.text)
    return " ".join(texts).strip()


# ──────────────────────────────────────────────────────
# Helper: print a full table with merged-cell handling
# ──────────────────────────────────────────────────────
def print_full_table(table, idx, label=""):
    header = f"{'='*80}\nTABLE {idx}"
    if label:
        header += f" - {label}"
    print(header)
    print(f"  Rows: {len(table.rows)},  Cols: {len(table.columns)}")
    print('='*80)
    for r_idx, row in enumerate(table.rows):
        cells_text = []
        for c_idx, cell in enumerate(row.cells):
            txt = extract_cell_text(cell).replace('\n', ' | ')
            cells_text.append(txt if txt else "(empty)")
        print(f"  Row {r_idx:3d}: " + "  ||  ".join(cells_text))
    print()


# ──────────────────────────────────────────────────────
# Helper: print data table with summary statistics
# ──────────────────────────────────────────────────────
def print_data_table_summary(table, idx, label=""):
    header = f"{'='*80}\nTABLE {idx}"
    if label:
        header += f" - {label}"
    print(header)
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    print(f"  Rows: {num_rows},  Cols: {num_cols}")
    print('='*80)

    all_rows_text = []
    for row in table.rows:
        cells_text = []
        for cell in row.cells:
            txt = extract_cell_text(cell).replace('\n', ' | ')
            cells_text.append(txt if txt else "(empty)")
        all_rows_text.append(cells_text)

    if not all_rows_text:
        print("  (empty table)")
        print()
        return

    # Header row
    print(f"  HEADER (Row 0): " + "  ||  ".join(all_rows_text[0]))

    data_rows = all_rows_text[1:]
    total_data = len(data_rows)
    print(f"  Total data rows (excluding header): {total_data}")

    # First 5 data rows
    print(f"\n  --- First 5 data rows ---")
    for i, row_text in enumerate(data_rows[:5]):
        print(f"  Row {i+1:3d}: " + "  ||  ".join(row_text))

    # Last 3 data rows
    if total_data > 5:
        print(f"\n  --- Last 3 data rows ---")
        for i, row_text in enumerate(data_rows[-3:]):
            actual_idx = total_data - 3 + i + 1
            print(f"  Row {actual_idx:3d}: " + "  ||  ".join(row_text))

    # Try to compute average of last column if numeric
    last_col_vals = []
    for row_text in data_rows:
        if row_text:
            val_str = row_text[-1].strip().replace(',', '.')
            # Try to parse as float
            try:
                last_col_vals.append(float(val_str))
            except ValueError:
                # Try extracting a number from the string
                m = re.search(r'[\d]+[.,]?\d*', val_str)
                if m:
                    try:
                        last_col_vals.append(float(m.group().replace(',', '.')))
                    except ValueError:
                        pass

    if last_col_vals:
        avg = statistics.mean(last_col_vals)
        print(f"\n  >> Last column ('{all_rows_text[0][-1]}') numeric values found: {len(last_col_vals)}")
        print(f"  >> Average of last column: {avg:.4f}")
        print(f"  >> Min: {min(last_col_vals):.4f},  Max: {max(last_col_vals):.4f}")
        if len(last_col_vals) > 1:
            print(f"  >> StdDev: {statistics.stdev(last_col_vals):.4f}")
    else:
        print(f"\n  >> No numeric values found in last column.")

    print()


# ──────────────────────────────────────────────────────
# Main extraction
# ──────────────────────────────────────────────────────
def main():
    print(f"Loading {DOCX_PATH} ...")
    doc = Document(DOCX_PATH)
    tables = doc.tables
    total_tables = len(tables)
    print(f"Total tables in document: {total_tables}\n")

    # ── 1) TABLE 32: Matriz de consistencia ──
    print("\n" + "#"*80)
    print("# SECTION 1: TABLE 32 - Matriz de consistencia")
    print("#"*80 + "\n")
    if total_tables > 32:
        print_full_table(tables[32], 32, "Matriz de consistencia")
    else:
        print(f"  Table 32 not found (only {total_tables} tables)")

    # ── 2) TABLES 33-38: Operationalization & Ficha de registro ──
    print("\n" + "#"*80)
    print("# SECTION 2: TABLES 33-38 - Operationalization / Ficha de registro")
    print("#"*80 + "\n")
    for i in range(33, 39):
        if i < total_tables:
            print_full_table(tables[i], i, "Operationalization / Ficha de registro")
        else:
            print(f"  Table {i} not found (only {total_tables} tables)")

    # ── 3) TABLES 40-45: Pre-test and Post-test data ──
    print("\n" + "#"*80)
    print("# SECTION 3: TABLES 40-45 - Pre-test / Post-test data")
    print("#"*80 + "\n")
    for i in range(40, 46):
        if i < total_tables:
            print_data_table_summary(tables[i], i, "Pre/Post test data")
        else:
            print(f"  Table {i} not found (only {total_tables} tables)")

    # ── 4) Search paragraphs for sampling methodology ──
    print("\n" + "#"*80)
    print("# SECTION 4: Paragraphs about sampling methodology")
    print("#"*80 + "\n")
    keywords = [
        r'poblaci[oó]n',
        r'muestra',
        r'N\s*=',
        r'n\s*=',
        r'\b53\b',
        r'\b50\b',
        r'muestreo',
        r'censo',
        r'inclusi[oó]n',
        r'exclusi[oó]n',
    ]
    pattern = re.compile('|'.join(keywords), re.IGNORECASE)

    matches_found = 0
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and pattern.search(text):
            matches_found += 1
            style = para.style.name if para.style else "None"
            print(f"  [Para {p_idx:4d}] (style={style})")
            # Print at most 500 chars to keep it readable
            display = text[:500]
            if len(text) > 500:
                display += " ..."
            print(f"    {display}")
            print()

    if matches_found == 0:
        print("  No paragraphs matched the sampling keywords.")
    else:
        print(f"  Total matching paragraphs: {matches_found}")

    # ── 5) Quick overview of nearby tables (30-46) to help identify contents ──
    print("\n" + "#"*80)
    print("# SECTION 5: Quick overview of tables 28-48 (first row only)")
    print("#"*80 + "\n")
    for i in range(28, min(49, total_tables)):
        row0 = tables[i].rows[0] if tables[i].rows else None
        if row0:
            cells = []
            for c in row0.cells:
                t = extract_cell_text(c).replace('\n', ' | ')[:80]
                cells.append(t if t else "(empty)")
            print(f"  Table {i:3d} ({len(tables[i].rows)} rows x {len(tables[i].columns)} cols): "
                  f"{' || '.join(cells)[:200]}")
    print()


if __name__ == "__main__":
    main()
